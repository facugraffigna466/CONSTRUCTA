#!/usr/bin/env python3
"""Genera docs/IPI-CONSTRUCTA.docx a partir de docs/IPI-CONSTRUCTA.md.

Conversor Markdown -> docx a medida para el IPI: respeta el tono académico
(texto justificado, encabezados, tablas, blockquotes resaltados para los
bloques [COMPLETAR]) e inserta los diagramas en las Figuras 1, 6 y 7.

Uso:  backend/.venv/bin/python docs/build_ipi_docx.py
"""
import os
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

DOCS = os.path.dirname(os.path.abspath(__file__))
MD = os.path.join(DOCS, "IPI-CONSTRUCTA.md")
OUT = os.path.join(DOCS, "IPI-CONSTRUCTA.docx")
DIAG = os.path.join(DOCS, "diagramas")

# Paleta CONSTRUCTA
INK = RGBColor(0x1A, 0x23, 0x29)
SLATE = RGBColor(0x3E, 0x4A, 0x52)
GREY = RGBColor(0x5B, 0x67, 0x70)
ORANGE = RGBColor(0xFF, 0x6B, 0x35)
BODY_FONT = "Calibri"

# Figura -> (archivo PNG, ancho en pulgadas). Las capturas (2-5) son marcadores.
FIG_IMG = {
    1: ("casos-de-uso.png", 5.4),
    6: ("arquitectura.png", 5.8),
    7: ("der.png", 6.0),
}


def shade(el, fill):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), fill)
    el.append(sh)


def left_border(par, color="FF6B35", size=18):
    pPr = par._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    pbdr.append(left)
    pPr.append(pbdr)


INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")


def add_runs(par, text, base_size=11, base_color=INK, base_italic=False):
    """Agrega runs interpretando **negrita**, *cursiva* y `monospace`."""
    for tok in INLINE.split(text):
        if not tok:
            continue
        bold = italic = mono = False
        t = tok
        if tok.startswith("**") and tok.endswith("**"):
            bold, t = True, tok[2:-2]
        elif tok.startswith("*") and tok.endswith("*"):
            italic, t = True, tok[1:-1]
        elif tok.startswith("`") and tok.endswith("`"):
            mono, t = True, tok[1:-1]
        r = par.add_run(t)
        r.font.size = Pt(base_size)
        r.font.color.rgb = base_color
        r.font.name = "Consolas" if mono else BODY_FONT
        r.bold = bold
        r.italic = italic or base_italic
        if mono:
            r.font.size = Pt(base_size - 0.5)


def body_par(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    add_runs(p, text)
    return p


def quote_par(doc, text):
    completar = "[COMPLETAR]" in text
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    left_border(p, "FF6B35" if completar else "C9D0D5")
    if completar:
        shade(p._p.get_or_add_pPr(), "FFF1E9")
    add_runs(p, text, base_size=10.5, base_color=SLATE, base_italic=not completar)
    return p


def add_figure(doc, num, desc):
    img = FIG_IMG.get(num)
    if img:
        fn, width = img
        path = os.path.join(DIAG, fn)
        pic_p = doc.add_paragraph()
        pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pic_p.add_run()
        run.add_picture(path, width=Inches(width))
        # alt-text en docPr
        try:
            inline = run._r.xpath(".//wp:inline")[0]
            docPr = inline.xpath(".//wp:docPr")[0]
            docPr.set("descr", desc)
        except Exception:
            pass
    else:
        # Marcador para captura de pantalla a insertar por el equipo
        ph = doc.add_paragraph()
        ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ph.paragraph_format.space_before = Pt(6)
        ph.paragraph_format.space_after = Pt(2)
        pPr = ph._p.get_or_add_pPr()
        shade(pPr, "F2F3F2")
        r = ph.add_run("[ Espacio reservado para captura de pantalla ]")
        r.italic = True
        r.font.size = Pt(10)
        r.font.color.rgb = GREY
        r.font.name = BODY_FONT
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    src = "elaboración propia"
    r = cap.add_run(f"Figura {num}. {desc} Fuente: {src}.")
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = GREY
    r.font.name = BODY_FONT


def add_heading(doc, text, level):
    sizes = {1: 16, 2: 13, 3: 11.5}
    colors = {1: INK, 2: INK, 3: SLATE}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(sizes[level])
    r.font.color.rgb = colors[level]
    r.font.name = BODY_FONT
    if level == 1:
        # regla naranja inferior
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "FF6B35")
        pbdr.append(bottom)
        pPr.append(pbdr)


def add_table(doc, rows):
    header, body = rows[0], rows[2:]  # rows[1] es el separador ---
    t = doc.add_table(rows=1, cols=len(header))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, cell in enumerate(header):
        c = t.rows[0].cells[i]
        c.paragraphs[0].text = ""
        add_runs(c.paragraphs[0], cell.strip(), base_size=10, base_color=RGBColor(0xFF, 0xFF, 0xFF))
        shade(c._tc.get_or_add_tcPr(), "1A2329")
    for ri, row in enumerate(body):
        cells = t.add_row().cells
        for i, cell in enumerate(row):
            if i < len(cells):
                cells[i].paragraphs[0].text = ""
                add_runs(cells[i].paragraphs[0], cell.strip(), base_size=10, base_color=INK)
                if ri % 2 == 1:
                    shade(cells[i]._tc.get_or_add_tcPr(), "F6F7F6")


def parse_table_row(line):
    return [c for c in line.strip().strip("|").split("|")]


def build_title_page(doc):
    def center(text, size, bold=False, color=INK, after=4, italic=False, font=BODY_FONT):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.color.rgb = color
        r.font.name = font
        return p

    doc.add_paragraph().paragraph_format.space_after = Pt(36)
    center("Universidad Católica de Córdoba", 15, bold=True)
    center("Facultad de Ingeniería", 13, color=SLATE, after=40)
    center("CONSTRUCTA", 30, bold=True, color=ORANGE, after=2)
    center("Plataforma de gestión de obras de construcción", 14, color=INK, after=2)
    center("con asistente de WhatsApp e inteligencia artificial", 14, color=INK, after=40)
    center("Informe de Proyecto Integrador", 15, bold=True, color=INK, after=36)

    center("Alumnos", 11, bold=True, color=SLATE, after=2)
    for _ in range(3):
        center("«Apellido, Nombre»", 11, color=INK, after=2)
    doc.add_paragraph().paragraph_format.space_after = Pt(14)
    center("Directores", 11, bold=True, color=SLATE, after=2)
    for _ in range(2):
        center("«Apellido, Nombre»", 11, color=INK, after=2)
    doc.add_paragraph().paragraph_format.space_after = Pt(40)
    center("«día» de «mes» de 2026", 11, color=GREY, after=2)
    center("Córdoba — Argentina", 11, color=GREY, after=2)
    doc.add_page_break()


def main():
    doc = Document()
    # Márgenes y estilo Normal
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(1)
        s.left_margin = s.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK

    build_title_page(doc)

    with open(MD, encoding="utf-8") as f:
        lines = f.read().split("\n")

    skip_portada = False
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Saltar la sección Portada (la portada se arma a mano arriba)
        if stripped == "## Portada":
            skip_portada = True
            i += 1
            continue
        if skip_portada:
            if stripped.startswith("## ") and stripped != "## Portada":
                skip_portada = False
            else:
                i += 1
                continue

        # Título H1 -> omitido (cubierto por portada)
        if stripped.startswith("# ") and not stripped.startswith("## "):
            i += 1
            continue

        if stripped == "---" or stripped == "":
            i += 1
            continue

        # Encabezados
        if stripped.startswith("#### "):
            add_heading(doc, stripped[5:].strip(), 3)
            i += 1
            continue
        if stripped.startswith("### "):
            add_heading(doc, stripped[4:].strip(), 2)
            i += 1
            continue
        if stripped.startswith("## "):
            add_heading(doc, stripped[3:].strip(), 1)
            i += 1
            continue

        # Tablas
        if stripped.startswith("|"):
            block = []
            while i < n and lines[i].strip().startswith("|"):
                block.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, block)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            continue

        # Blockquotes (incluye FIGURA y [COMPLETAR])
        if stripped.startswith(">"):
            content = stripped[1:].strip()
            fig = re.search(r"\[FIGURA\s+(\d+)\s*:\s*(.+?)\]", content)
            if fig:
                num = int(fig.group(1))
                desc = fig.group(2).strip()
                if not desc.endswith("."):
                    desc += "."
                add_figure(doc, num, desc)
            elif content:
                quote_par(doc, content)
            i += 1
            continue

        # Listas con viñeta
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(3)
            add_runs(p, stripped[2:].strip())
            i += 1
            continue

        # Listas numeradas
        m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(3)
            add_runs(p, m.group(2).strip())
            i += 1
            continue

        # Párrafo normal
        body_par(doc, stripped)
        i += 1

    doc.save(OUT)
    print(f"OK -> {OUT}")


if __name__ == "__main__":
    main()
